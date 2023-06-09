from requests import request
from rest_framework import generics, authentication, permissions, status
from rest_framework.response import Response
from .permissions import IsStaff
from .models import *
from .serializers import *
from django.http import HttpResponse, JsonResponse
from rest_framework.views import APIView
from rest_framework.serializers import ValidationError
from rest_framework.response import Response
from rest_framework.permissions import AllowAny, IsAuthenticated
from rest_framework.generics import UpdateAPIView
from rest_framework import generics, status
from django_filters.rest_framework import DjangoFilterBackend
import datetime
from django.conf import settings
from docx import Document
import matplotlib.pyplot as plt
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity


class registerUser(APIView):
    permission_classes = [AllowAny]

    def post(self, request, format=None):
        data = request.data
        user = User.objects.create_user(
            data["username"], data["email"], data["password"])
        user.name = data["name"]
        user.is_staff = data["is_staff"]
        user.phone = data["phone"]
        user.class_num = data["class_num"]
        user.save()
        serializer = UserSerializer(user)
        return Response(serializer.data)


class UsersView(generics.RetrieveAPIView):
    def get_permissions(self):
        method = self.request.method
        if method == 'GET':
            permission_classes = [AllowAny]
        else:
            permission_classes = [IsAuthenticated]
        return [permission() for permission in permission_classes]

    queryset = User.objects.all()
    serializer_class = UserSerializer


class ProfileView(generics.RetrieveUpdateDestroyAPIView):
    permission_classes = [IsAuthenticated]
    serializer_class = UserSerializer

    def get_object(self):
        return self.request.user

    def update(self, request, *args, **kwargs):
        partial = kwargs.pop('partial', False)
        serializer = UserUpdateSerializer(
            instance=request.user, data=request.data, partial=partial)
        serializer.is_valid(raise_exception=True)
        serializer.save()

        return Response(serializer.data, status=status.HTTP_200_OK)

    def destroy(self, request, *args, **kwargs):
        request.user.delete()
        return Response(status=status.HTTP_204_NO_CONTENT)


class ChangePasswordView(UpdateAPIView):
    permission_classes = [IsAuthenticated]
    serializer_class = ChangePasswordSerializer

    def update(self, request, *args, **kwargs):
        serializer = self.get_serializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        serializer.save()
        return Response(status=status.HTTP_200_OK)


class Submit_paper(generics.ListCreateAPIView):
    permission_classes = [permissions.IsAuthenticated]
    queryset = Paper.objects.all()
    serializer_class = PaperSerializer

    def get_queryset(self):
        return Paper.objects.filter(user=self.request.user)

    def perform_create(self, serializer):
        serializer.save(user=self.request.user)


class Show_paper(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request):
        if self.request.user.is_staff == True:
            class_num = request.query_params.get('class_num')
            users = User.objects.filter(class_num=class_num)
            data = Paper.objects.filter(user__in=users)
            serializer = PaperSerializer(data, many=True)
            return Response(serializer.data)
        else:
            return Response(status=status.HTTP_400_BAD_REQUEST)


class Check_result(APIView):
    permission_classes = [AllowAny]

    def get(self, request):
        if self.request.user.is_staff == True:
            class_num = request.query_params.get('class_num')

            if not class_num:
                return Response([])
            users = User.objects.filter(class_num=class_num, is_staff=False)

            data = Paper.objects.filter(user__in=users, user__is_staff=False)
            filenames = [p.paper.name for p in data[::-1]]
            plag = []
            result = check_plagiarism(filenames, plag)
            return JsonResponse(result, safe=False)
        else:
            return JsonResponse(status=status.HTTP_400_BAD_REQUEST)


class Check_answers(APIView):
    permission_classes = [AllowAny]

    def get(self, request):
        if self.request.user.is_staff == True:
            class_num = self.request.user.class_num
            users = User.objects.filter(class_num=class_num, is_staff=False)

            data1 = Paper.objects.filter(user__in=users, user__is_staff=False)
            data2 = Paper.objects.filter(user=self.request.user)
            filenames = [p.paper.name for p in data1[::-1]]
            filename = [p.paper.name for p in data2[::-1]]
            plag = []
            result = score_answers(filenames, filename)
            return JsonResponse(result, safe=False)
        else:
            return JsonResponse(status=status.HTTP_400_BAD_REQUEST)


def check_plagiarism(data, plag):

    path = '''C:/Users/Rog/OneDrive/Desktop/ml practical/project_api/plagiarism_classroom/paper/'''
    doc_files = data
    num_docs = len(doc_files)
    if num_docs < 2:
        return "Not enough documents to compare for plagiarism."

    docs = []
    for file in doc_files:
        f_name = file.split('/')[-1]
        document = Document(path+f_name)
        full_text = []
        for para in document.paragraphs:
            full_text.append(para.text)
        doc_text = '\n'.join(full_text)
        docs.append(doc_text)

    vectorizer = TfidfVectorizer(stop_words='english')
    doc_vectors = vectorizer.fit_transform(docs)
    similarities = cosine_similarity(doc_vectors)

    results = []

    for i in range(num_docs):
        count = 0
        for j in range(num_docs):
            if i == j:
                continue
            similarity = similarities[i][j]

            if similarity > 0.8:
                count += similarity
                score = (num_docs-i)/similarity
                result = f"Plagiarism,{doc_files[i]},{doc_files[j]},{similarity:.2f},{score:.2f}"
            else:
                score = (num_docs-i)/similarity
                result = f"No,{doc_files[i]},{doc_files[j]},{similarity:.2f},{score:.2f}"
            results.append(result)
        plag.append(count)

    return results


def score_answers(data1, data2):

    path = '''C:/Users/Rog/OneDrive/Desktop/ml practical/project_api/plagiarism_classroom/paper/'''
    doc_files = data1 + data2
    num_docs = len(doc_files)
    if num_docs < 2:
        return "Not enough documents to compare for plagiarism."

    docs = []

    for file in doc_files:
        f_name = file.split('/')[-1]
        document = Document(path+f_name)
        full_text = []
        for para in document.paragraphs:
            full_text.append(para.text)
        doc_text = '\n'.join(full_text)
        docs.append(doc_text)

    vectorizer = TfidfVectorizer(stop_words='english')
    doc_vectors = vectorizer.fit_transform(docs)
    similarities = cosine_similarity(doc_vectors)

    results = []

    for j in range(num_docs-1):
        similarity = similarities[j][num_docs-1]
        result = f"score,{doc_files[j]},{doc_files[num_docs-1]},{similarity:.2f}"
        results.append(result)

    return results
